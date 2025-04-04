from pathlib import Path
from docx import Document
from fpdf import FPDF

def create_txt(filename: str):
    content = """\
¿Cuáles son los horarios de atención?
Nuestro horario es de lunes a viernes de 8:00 a.m. a 6:00 p.m.

¿Ofrecen soporte remoto?
Sí. Nuestro equipo de soporte está capacitado para atenderte de forma remota usando herramientas seguras.

¿Puedo solicitar mantenimiento presencial?
Sí. Tenemos técnicos disponibles para visitas programadas en todo el país.
"""
    Path(filename).write_text(content, encoding="utf-8")
    print(f"✅ Archivo creado: {filename}")


def create_docx(filename: str):
    doc = Document()
    doc.add_heading("Servicios de NovaTech Solutions", level=1)
    doc.add_paragraph("1. Desarrollo de sitios web corporativos, eCommerce y landing pages.")
    doc.add_paragraph("2. Automatización de procesos empresariales.")
    doc.add_paragraph("3. Instalación y configuración de software empresarial.")
    doc.add_paragraph("4. Consultoría en ciberseguridad y respaldo en la nube.")
    doc.add_paragraph("5. Venta de licencias de software y hardware especializado.")
    doc.save(filename)
    print(f"✅ Archivo creado: {filename}")


def create_pdf(filename: str):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, 
        "Política de devoluciones\n"
        "Los productos físicos tienen un plazo de 10 días calendario para solicitar devolución.\n"
        "El producto debe estar sin uso y en su empaque original.\n\n"
        "Política de privacidad\n"
        "Respetamos tu privacidad. No compartimos datos con terceros sin tu consentimiento.\n"
        "Garantía de servicios\n"
        "Todos nuestros servicios tienen garantía de satisfacción por 30 días."
    )
    pdf.output(filename)
    print(f"✅ Archivo creado: {filename}")


if __name__ == "__main__":
    create_txt("faq.txt")
    create_docx("servicios.docx")
    create_pdf("politicas.pdf")
    print("🎉 Todos los archivos se han generado con éxito.")
